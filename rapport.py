from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn



def creer_rapport():
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.5

    # ---- PAGE DE GARDE ----
    for _ in range(6):
        doc.add_paragraph()

    titre = doc.add_paragraph()
    titre.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = titre.add_run("RAPPORT DE PROJET")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x7C, 0x3A, 0xED)

    sous_titre = doc.add_paragraph()
    sous_titre.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sous_titre.add_run("Planification d'Examens par Coloration de Graphe")
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    doc.add_paragraph()

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run("Theorie des Graphes - L2 Informatique\nAnnee universitaire 2025-2026")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

    doc.add_page_break()

    # ---- SOMMAIRE ----
    doc.add_heading("Sommaire", level=1)

    sommaire = [
        "1. Introduction",
        "2. Objectifs du projet",
        "3. Architecture technique",
        "4. Modele de donnees",
        "5. Construction du graphe de conflits",
        "6. Algorithmes de coloration",
        "7. Affectation des salles et des creneaux",
        "8. Verification des contraintes",
        "9. Export CSV et PDF",
        "10. Securite multi-etablissements",
        "11. Interface responsive",
        "12. Guide d'utilisation",
        "13. Conclusion",
    ]

    for item in sommaire:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(4)

    doc.add_page_break()

    # ---- 1. INTRODUCTION ----
    doc.add_heading("1. Introduction", level=1)

    doc.add_paragraph(
        "Ce projet consiste a developper une application web de planification "
        "d'examens universitaires en utilisant les concepts de la theorie des graphes, "
        "notamment les algorithmes de coloration de graphe. L'objectif est d'automatiser "
        "l'organisation des examens en tenant compte des contraintes reelles : "
        "conflits d'horaires, disponibilite des salles, interdictions de surveillance, etc."
    )

    doc.add_paragraph(
        "L'application propose une interface web complete permettant de saisir les donnees, "
        "visualiser le graphe de conflits, lancer les algorithmes de coloration, et generer "
        "un planning optimise exportable en CSV et PDF."
    )

    # ---- 2. OBJECTIFS ----
    doc.add_heading("2. Objectifs du projet", level=1)

    objectifs = [
        "Modeliser le probleme de planification d'examens sous forme de graphe de conflits",
        "Implementer les algorithmes Welsh-Powell et DSATUR pour la coloration",
        "Automatiser l'affectation des salles et des creneaux",
        "Gerer les contraintes specifiques (filiere eloignee, charge equilibree)",
        "Proposer une interface web intuitive et responsive",
        "Permettre l'export du planning en CSV et PDF",
        "Isoler les donnees par etablissement avec authentification par code",
    ]

    for obj in objectifs:
        doc.add_paragraph(obj, style='List Bullet')

    # ---- 3. ARCHITECTURE ----
    doc.add_heading("3. Architecture technique", level=1)

    doc.add_paragraph(
        "L'application est developpee en Python avec le framework Flask. "
        "La base de donnees utilise SQLite via SQLAlchemy. La visualisation "
        "des graphes est realisee avec Matplotlib et NetworkX."
    )

    doc.add_heading("Stack technique", level=3)

    table = doc.add_table(rows=8, cols=2)
    table.style = 'Light Shading Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    donnees = [
        ("Composant", "Technologie"),
        ("Backend", "Python / Flask"),
        ("Base de donnees", "SQLite / SQLAlchemy"),
        ("Graphes", "NetworkX"),
        ("Visualisation", "Matplotlib"),
        ("Export PDF", "ReportLab"),
        ("Frontend", "HTML5 / CSS3 / JavaScript"),
        ("Polices", "Google Fonts (Inter)"),
    ]

    for i, (col1, col2) in enumerate(donnees):
        table.rows[i].cells[0].text = col1
        table.rows[i].cells[1].text = col2
        if i == 0:
            for cell in table.rows[i].cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

    doc.add_paragraph()

    doc.add_heading("Structure des fichiers", level=3)

    structure = [
        "app.py - Point d'entree, configuration Flask",
        "config.py - Configuration de la base de donnees",
        "extensions.py - Initialisation de SQLAlchemy",
        "models/ - Modeles de donnees (UE, Salle, Etudiant, Conflit, etc.)",
        "routes/ - Routes Flask (Blueprints) pour chaque section",
        "algorithms/ - Implementation de Welsh-Powell et DSATUR",
        "graph/ - Construction du graphe, matrice et liste d'adjacence",
        "scheduling/ - Affectation, verification, export CSV/PDF",
        "templates/ - Templates HTML (Jinja2)",
        "static/css/ - Feuille de style CSS",
    ]

    for item in structure:
        doc.add_paragraph(item, style='List Bullet')

    # ---- 4. MODELE DE DONNEES ----
    doc.add_heading("4. Modele de donnees", level=1)

    doc.add_paragraph(
        "Le modele de donnees est organise autour de plusieurs entites principales, "
        "toutes liees a un etablissement pour l'isolation des donnees."
    )

    table = doc.add_table(rows=9, cols=3)
    table.style = 'Light Shading Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    modeles = [
        ("Modele", "Champs principaux", "Description"),
        ("Etablissement", "nom, adresse, code", "Unite d'organisation isolee"),
        ("UE", "code, nom, filiere, professeur", "Unite d'enseignement"),
        ("Etudiant", "nom, prenom, matricule", "Etudiant inscrit"),
        ("Inscription", "etudiant_id, ue_id", "Lien etudiant-UE"),
        ("Salle", "nom, capacite", "Salle d'examen"),
        ("Creneau", "jour, heure_debut, heure_fin", "Plage horaire"),
        ("Conflit", "ue1_id, ue2_id", "Conflit Manuel"),
        ("Interdiction", "ue1_id, ue2_id", "Interdiction explicite"),
    ]

    for i, (c1, c2, c3) in enumerate(modeles):
        table.rows[i].cells[0].text = c1
        table.rows[i].cells[1].text = c2
        table.rows[i].cells[2].text = c3
        if i == 0:
            for cell in table.rows[i].cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

    # ---- 5. CONSTRUCTION DU GRAPHE ----
    doc.add_heading("5. Construction du graphe de conflits", level=1)

    doc.add_paragraph(
        "Le graphe de conflits est construit automatiquement a partir des donnees saisies. "
        "Chaque UE devient un noeud du graphe. Une arete est ajoutee entre deux UE si "
        "elles ne peuvent pas avoir lieu au meme creneau."
    )

    doc.add_heading("Types de conflits consideres", level=3)

    conflits = [
        "Conflits manuels : Ajoutes directement par l'utilisateur",
        "Etudiants communs : Deux UE partageant au moins un etudiant",
        "Meme professeur : Un prof ne peut pas evaluer deux UE au meme moment",
        "Interdictions explicites : Regles specifiques definies par l'utilisateur",
    ]

    for c in conflits:
        doc.add_paragraph(c, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph(
        "La construction est implementee dans graph/graphe.py via la fonction "
        "construire_graphe(etablissement_id). Le graphe est represente avec "
        "NetworkX (nx.Graph)."
    )

    doc.add_heading("Representation du graphe", level=3)

    doc.add_paragraph(
        "Deux representations sont disponibles :"
    )

    rep = [
        "Matrice d'adjacence : Tableau NxN avec 1 si conflit, 0 sinon",
        "Liste d'adjacence : Pour chaque noeud, la liste de ses voisins",
    ]

    for r in rep:
        doc.add_paragraph(r, style='List Bullet')

    # ---- 6. ALGORITHMES ----
    doc.add_heading("6. Algorithmes de coloration", level=1)

    doc.add_paragraph(
        "La coloration de graphe est au coeur du projet. Chaque couleur correspond "
        "a un creneau horaire different. L'objectif est de minimiser le nombre de "
        "couleurs (creneaux) utilisees."
    )

    doc.add_heading("6.1 Welsh-Powell", level=2)

    doc.add_paragraph(
        "Algorithme glouton (greedy) qui fonctionne ainsi :"
    )

    wp_etapes = [
        "Trier les sommets par degre decroissant (nombre de conflits)",
        "Pour chaque sommet, lui attribuer la plus petite couleur non utilisee par ses voisins",
        "Avancer au sommet suivant et repeter",
    ]

    for i, e in enumerate(wp_etapes, 1):
        doc.add_paragraph(f"{i}. {e}")

    doc.add_paragraph(
        "Dans notre implementation, les sommets sont egalement tries par effectif "
        "(nombre d'etudiants inscrits) pour privilegier les UE les plus chargees. "
        "Complexite : O(n^2)."
    )

    doc.add_heading("6.2 DSATUR", level=2)

    doc.add_paragraph(
        "Algorithme plus precis qui utilise la notion de saturation :"
    )

    ds_etapes = [
        "Initialiser la saturation de chaque noeud a 0",
        "Choisir le noeud non colore avec la plus haute saturation (et degre en cas d'egalite)",
        "Lui attribuer la plus petite couleur non utilisee par ses voisins",
        "Mettre a jour la saturation de ses voisins non colores",
        "Repeter jusqu'a ce que tous les noeuds soient colores",
    ]

    for i, e in enumerate(ds_etapes, 1):
        doc.add_paragraph(f"{i}. {e}")

    doc.add_paragraph(
        "DSATUR produit generalement une coloration optimale ou quasi-optimale. "
        "Complexite : O(n^3)."
    )

    doc.add_heading("6.3 Comparaison", level=2)

    table = doc.add_table(rows=4, cols=3)
    table.style = 'Light Shading Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    comp = [
        ("Critere", "Welsh-Powell", "DSATUR"),
        ("Complexite", "O(n^2)", "O(n^3)"),
        ("Qualite", "Bonne", "Excellent"),
        ("Vitesse", "Rapide", "Plus lent"),
    ]

    for i, (c1, c2, c3) in enumerate(comp):
        table.rows[i].cells[0].text = c1
        table.rows[i].cells[1].text = c2
        table.rows[i].cells[2].text = c3
        if i == 0:
            for cell in table.rows[i].cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

    # ---- 7. AFFECTATION ----
    doc.add_heading("7. Affectation des salles et des creneaux", level=1)

    doc.add_paragraph(
        "Apres la coloration, le systeme affecte automatiquement les salles "
        "et les creneaux reels a chaque UE."
    )

    doc.add_heading("7.1 Mapping creneau index / creneau reel", level=3)

    doc.add_paragraph(
        "Chaque couleur (= index de creneau) est associee a un creneau reel "
        "(jour + heure_debut + heure_fin) depuis la table des creneaux."
    )

    doc.add_heading("7.2 Affectation des salles", level=3)

    doc.add_paragraph(
        "Les salles sont assignees en fonction de la capacite : "
        "le systeme choisit la plus petite salle dont la capacite est "
        "suffisante pour accueillir les etudiants de l'UE."
    )

    doc.add_heading("7.3 Charge par creneau", level=3)

    doc.add_paragraph(
        "Le systeme calcule le nombre d'examens simultanes a chaque creneau "
        "et affiche une alerte si la charge est desequilibree (trop d'examens "
        "au meme creneau)."
    )

    # ---- 8. VERIFICATION ----
    doc.add_heading("8. Verification des contraintes", level=1)

    doc.add_paragraph(
        "Avant de valider le planning, le systeme verifie plusieurs contraintes :"
    )

    contraintes = [
        "Aucun conflit violee : Deux UE en conflit ne sont pas au meme creneau",
        "Filiere eloignee : Deux UE de memes filieres sont placees a des creneaux eloignes",
        "Charge equilibree : Le nombre d'examens par creneau est reparti uniformement",
        "Interdictions respectees : Les regles de surveillance sont respectees",
    ]

    for c in contraintes:
        doc.add_paragraph(c, style='List Bullet')

    doc.add_paragraph(
        "Des alertes sont generees en cas de violation, permettant a l'utilisateur "
        "de relancer la generation pour obtenir un meilleur resultat."
    )

    # ---- 9. EXPORT ----
    doc.add_heading("9. Export CSV et PDF", level=1)

    doc.add_heading("9.1 Export CSV", level=2)

    doc.add_paragraph(
        "Le fichier CSV est genere au format Creneau x Salle, chaque cellule "
        "contenant le code de l'UE et son effectif. Le format est compatible "
        "avec Excel et Google Sheets."
    )

    doc.add_heading("9.2 Export PDF", level=2)

    doc.add_paragraph(
        "Le PDF est genere avec ReportLab en format A4 paysage. "
        "Le tableau utilise un schema de couleurs clair (fond blanc, "
        "en-tetes violets, lignes alternees gris tres clair) pour "
        "une impression optimale."
    )

    # ---- 10. SECURITE ----
    doc.add_heading("10. Securite multi-etablissements", level=1)

    doc.add_paragraph(
        "L'application supporte plusieurs etablissements de maniere isolee. "
        "Chaque etablissement est protege par un code de 6 caracteres genere "
        "automatiquement a sa creation."
    )

    doc.add_paragraph(
        "Pour activer un etablissement, l'utilisateur doit selectionner "
        "l'etablissement dans la liste et entrer le code correspondant. "
        "Sans ce code, l'etablissement ne peut pas etre actif et les "
        "donnees ne sont pas accessibles."
    )

    doc.add_paragraph(
        "L'isolation est assuree par le stockage de l'etablissement actif "
        "dans la session Flask, et tous les requetes de donnees filtrent "
        "par etablissement_id."
    )

    # ---- 11. RESPONSIVE ----
    doc.add_heading("11. Interface responsive", level=1)

    doc.add_paragraph(
        "L'interface est entierement responsive grace a des media queries CSS. "
        "Le site s'adapte a toutes les tailles d'ecran :"
    )

    responsive = [
        "Desktop (> 900px) : Sidebar fixe a gauche, contenu principal avec marge",
        "Tablette (<= 900px) : Sidebar masquee, bouton hamburger pour l'ouvrir",
        "Mobile (<= 600px) : Mise en page compacte, cartes en grille 2 colonnes",
    ]

    for r in responsive:
        doc.add_paragraph(r, style='List Bullet')

    # ---- 12. GUIDE ----
    doc.add_heading("12. Guide d'utilisation rapide", level=1)

    etapes = [
        ("Creer un etablissement", "Nommer l'etablissement, noter le code genere, le selectionner avec le code"),
        ("Saisir les donnees", "Ajouter les UE, etudiants, inscriptions, salles et creneaux"),
        ("Definir les contraintes", "Ajouter les conflits et interdictions specifiques"),
        ("Visualiser le graphe", "Consulter le graphe de conflits, la matrice et la liste d'adjacence"),
        ("Comparer les colorations", "Lancer Welsh-Powell et DSATUR, comparer le nombre de couleurs"),
        ("Generer le planning", "Cliquez sur Generer, verifier les alertes, relancer si necessaire"),
        ("Telecharger", "Exporter en CSV ou PDF depuis la page Planning"),
    ]

    for i, (titre, desc) in enumerate(etapes, 1):
        doc.add_paragraph(f"Etape {i} - {titre}", style='List Number')
        doc.add_paragraph(desc)

    # ---- 13. CONCLUSION ----
    doc.add_heading("13. Conclusion", level=1)

    doc.add_paragraph(
        "Ce projet permet de resoudre concrètement le probleme de la planification "
        "d'examens en utilisant les algorithmes de coloration de graphe. "
        "L'application web developpee offre une solution complete, de la saisie "
        "des donnees a l'export du planning final."
    )

    doc.add_paragraph(
        "Les algorithmes Welsh-Powell et DSATUR sont compares et permettent "
        "d'obtenir un planning optimise, sans conflits, avec une affectation "
        "automatique des salles et des creneaux. L'interface responsive et "
        "le systeme de securite multi-etablissements rendent l'application "
        "utilisable en conditions reelles."
    )

    # Sauvegarde
    chemin = "static/rapport_projet.docx"
    import os
    os.makedirs("static", exist_ok=True)
    doc.save(chemin)

    return chemin


if __name__ == "__main__":
    chemin = creer_rapport()
    print(f"Rapport genere : {chemin}")
